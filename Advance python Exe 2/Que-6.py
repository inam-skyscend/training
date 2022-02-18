# Create a word document with header as Employee Experience Certificate, Write a
# paragraph for his experience description.
# Mention his projects as ordered list.
# Mention his skills with unordered list.
# Create a table with the following details.


import docx
from docx import Document
document=Document()

#add heading
document.add_heading('Employee Experience Certificate').alignment=1
# add heading to center  alignment 1 for center
# alignment=1

document.add_paragraph('\nI have 4 years experience in Python')
document.add_paragraph('Projects:')


# projects in ordered list.
document.add_paragraph("Fraud email detection",style='List Number')
document.add_paragraph("Website development",style='List Number')
document.add_paragraph("Placement Management System in Android",style='List Number')


document.add_paragraph('Skills:').bold=True
document.add_paragraph("Python Programming",style='List Bullet')
document.add_paragraph("Web development",style='List Bullet')
document.add_paragraph("Travel",style='List Bullet')



table=document.add_table(4,4)

hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Total number of year'
hdr_cells[1].text = 'Total Working Days'
hdr_cells[2].text = 'Worked Days'
hdr_cells[3].text = 'Leaves'

hdr_cells = table.rows[1].cells
hdr_cells[0].text = "4"
hdr_cells[1].text = "1056"
hdr_cells[2].text = '1016'
hdr_cells[3].text = '40'


hdr_cells = table.rows[2].cells
hdr_cells[0].text = 'Total Leaves'
hdr_cells[1].text = 'Paid Leave'
hdr_cells[2].text = 'Sick Leave'
hdr_cells[3].text = 'Unpaid Leave'

hdr_cells = table.rows[3].cells
hdr_cells[0].text = "40"
hdr_cells[1].text = "25"
hdr_cells[2].text = '5'
hdr_cells[3].text = '10'


document.save("Certificate.docx")
