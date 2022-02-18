# Read the data written in the above document file. For table data prepare a
# dictionary with their headers.

from docx import Document
document=Document("Certificate.docx")

print(document.paragraphs)
paras = [p.text for p in document.paragraphs]
print(paras)

for table in document.tables:
    data = [[cell.text for cell in row.cells] for row in table.rows]


    print(dict(zip(data[0],data[1])))
    print(dict(zip(data[2],data[3])))







